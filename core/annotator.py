
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import os

def annotate_docx(path, issues):
    """Creates a copy of the docx and inserts visible annotations near the top and at document end."""
    doc = Document(path)
    # Insert summary at top
    if issues:
        doc_para = doc.add_paragraph()
        doc_para.add_run('--- ADGM Corporate Agent Review Annotations ---').bold = True
        for i, it in enumerate(issues, start=1):
            r = doc.add_paragraph()
            run = r.add_run(f"[{i}] {it.get('section')}: {it.get('issue')} Suggested: {it.get('suggestion')}")
            run.bold = False
            run.font.color.rgb = RGBColor(0x9B, 0x00, 0x00)
    else:
        doc.add_paragraph('No issues detected by heuristic checks.')
    # Save annotated file
    base, name = os.path.split(path)
    out = os.path.join(base, 'annotated_' + name)
    doc.save(out)
    return out
